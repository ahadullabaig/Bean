from docx import Document

def create_master_template():
    doc = Document()
    
    # Header / Title
    doc.add_heading('Event Report: {{ facts.event_title }}', 0)
    
    # Meta Data Table
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date: {{ facts.date }}'
    hdr_cells[1].text = 'Venue: {{ facts.venue }}'
    
    doc.add_paragraph(f"Attendance: {{{{ facts.attendance_count }}}}")
    doc.add_paragraph(f"Speaker: {{{{ facts.speaker_name }}}}")
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph('{{ executive_summary }}')
    
    doc.add_heading('Key Takeaways', level=1)
    # Loop for takeaways
    # NOTE: python-docx doesn't natively support jinja2 tags perfectly in all add_paragraph calls 
    # if we want complex structures, but for text it's fine.
    # For a loop, we textually write the Jinja2 loop syntax.
    
    doc.add_paragraph('{% for item in key_takeaways %}')
    doc.add_paragraph('- {{ item }}', style='List Bullet')
    doc.add_paragraph('{% endfor %}')
    
    doc.save('master_template.docx')
    print("master_template.docx created successfully.")

if __name__ == "__main__":
    create_master_template()
